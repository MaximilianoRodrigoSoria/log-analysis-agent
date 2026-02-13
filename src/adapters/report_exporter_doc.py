"""
Exporter de reportes en formato DOC/DOCX usando python-docx.
Genera documentos Word con formato profesional.
"""

import logging
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..ports.report_exporter_port import ReportExporterPort


logger = logging.getLogger(__name__)


class DocExporter(ReportExporterPort):
    """Exporta reportes en formato DOCX (Microsoft Word)"""
    
    def export(
        self,
        output_dir: str,
        output_filename: str,
        report_content: str,
        analysis: Optional[Dict] = None
    ) -> str:
        """
        Exporta el reporte en formato DOCX.
        Crea un documento Word con formato profesional.
        
        Args:
            output_dir: Directorio de salida
            output_filename: Nombre del archivo (sin extensión)
            report_content: Contenido del reporte en Markdown
            analysis: Análisis estructurado (para secciones adicionales)
        
        Returns:
            Path absoluto del archivo generado
        
        Raises:
            IOError: Si hay error de escritura
        """
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        filename = f"{output_filename}.docx"
        file_path = output_path / filename
        
        logger.debug(f"Exportando reporte DOCX a {file_path}")
        
        try:
            document = Document()
            
            # Configurar márgenes y estilos
            self._setup_document_style(document)
            
            # Título principal
            title = document.add_heading('Reporte de Análisis de Logs', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Separador
            document.add_paragraph()
            
            # Resumen ejecutivo si hay analysis
            if analysis and "summary" in analysis:
                self._add_summary_section(document, analysis["summary"])
            
            # Contenido del reporte (convertir markdown a párrafos)
            self._add_report_content(document, report_content)
            
            # Detalles de errores si hay analysis
            if analysis and "error_groups" in analysis:
                self._add_errors_section(document, analysis["error_groups"])
            
            # Advertencias si hay analysis
            if analysis and "warnings" in analysis:
                self._add_warnings_section(document, analysis["warnings"])
            
            # Pie de página
            document.add_page_break()
            footer_para = document.add_paragraph()
            footer_para.add_run('Generado automáticamente por Log Analyzer').italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Guardar documento
            document.save(str(file_path))
            
            logger.info(f"Reporte DOCX exportado: {file_path}")
            return str(file_path.absolute())
        
        except Exception as e:
            logger.error(f"Error al exportar reporte DOCX: {e}")
            raise IOError(f"Error al escribir archivo DOCX: {e}") from e
    
    def _setup_document_style(self, document: Document):
        """Configura estilos básicos del documento"""
        # Los estilos se heredan de las plantillas por defecto de python-docx
        # Aquí podríamos personalizar más si fuera necesario
        pass
    
    def _add_summary_section(self, document: Document, summary: Dict):
        """Agrega sección de resumen ejecutivo"""
        document.add_heading('Resumen Ejecutivo', level=1)
        
        # Crear tabla para el resumen
        table = document.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        # Datos
        for key, value in summary.items():
            row_cells = table.add_row().cells
            label = key.replace('_', ' ').title()
            row_cells[0].text = label
            row_cells[1].text = str(value)
        
        document.add_paragraph()
    
    def _add_report_content(self, document: Document, content: str):
        """Agrega el contenido principal del reporte"""
        document.add_heading('Análisis Detallado', level=1)
        
        # Convertir markdown básico a párrafos
        lines = content.split('\n')
        current_list = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                current_list = None
                continue
            
            # Encabezados
            if line.startswith('# '):
                document.add_heading(line[2:], level=1)
                current_list = None
            elif line.startswith('## '):
                document.add_heading(line[3:], level=2)
                current_list = None
            elif line.startswith('### '):
                document.add_heading(line[4:], level=3)
                current_list = None
            
            # Listas
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                # Remover formato markdown básico
                text = text.replace('**', '').replace('__', '').replace('`', '')
                document.add_paragraph(text, style='List Bullet')
                current_list = True
            
            # Párrafo normal
            else:
                # Remover formato markdown básico
                text = line.replace('**', '').replace('__', '').replace('`', '')
                if text:
                    p = document.add_paragraph(text)
                    current_list = None
    
    def _add_errors_section(self, document: Document, error_groups: list):
        """Agrega sección de errores detallados"""
        document.add_heading('Detalle de Errores', level=1)
        
        if not error_groups:
            document.add_paragraph('No se encontraron errores.')
            return
        
        # Crear tabla
        table = document.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tipo de Error'
        hdr_cells[1].text = 'Ocurrencias'
        hdr_cells[2].text = 'Componente'
        hdr_cells[3].text = 'Primera Aparición'
        
        # Datos (limitar a 20 errores)
        for error_group in error_groups[:20]:
            row_cells = table.add_row().cells
            row_cells[0].text = error_group.get('exception', 'Unknown')
            row_cells[1].text = str(error_group.get('count', 0))
            row_cells[2].text = error_group.get('logger', 'Unknown')
            row_cells[3].text = error_group.get('first_ts', 'N/A')
        
        document.add_paragraph()
    
    def _add_warnings_section(self, document: Document, warnings: list):
        """Agrega sección de advertencias"""
        document.add_heading('Advertencias', level=1)
        
        if not warnings:
            document.add_paragraph('No se encontraron advertencias.')
            return
        
        # Limitar a 10 warnings
        for warning in warnings[:10]:
            timestamp = warning.get('timestamp', 'N/A')
            component = warning.get('logger', 'Unknown')
            message = warning.get('message', '')
            
            p = document.add_paragraph()
            p.add_run(f'[{timestamp}] ').bold = True
            p.add_run(f'{component}: ')
            p.add_run(message)
        
        if len(warnings) > 10:
            document.add_paragraph(
                f'... y {len(warnings) - 10} advertencias adicionales.'
            ).italic = True
